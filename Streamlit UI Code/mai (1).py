import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
from datetime import datetime, timedelta
import time
import streamlit.components.v1 as components

# ======================================================
# PAGE CONFIG
# ======================================================
from Background import BackgroundCSSGenerator
st.set_page_config(
    page_title="Teacher Performance Tracker",
    page_icon="📊",  # You can replace this with an emoji or a URL to an image
    layout="wide"
)
# img1_path = r"200w.gif"



# background_generator = BackgroundCSSGenerator(img1_path)
# page_bg_img = background_generator.generate_background_css()
# st.markdown(page_bg_img, unsafe_allow_html=True)
# # Dummy user credentials
st.markdown("""
<style>
.metric-card {
    padding: 12px;
    border-radius: 14px;
    background: rgba(240,240,240,0.5);
}
.control-box {
    padding: 16px;
    border-radius: 16px;
    background: linear-gradient(145deg, #f7f9fc, #eef1f6);
}
</style>
""", unsafe_allow_html=True)

np.random.seed(42)

BATTERY_ID = "BAT-AADHAAR-IND-00042"
N_CELLS_SERIES = 3

import serial

try:
    ser = serial.Serial("COM9", 9600, timeout=1)
except:
    ser = None


def read_sensor_data():
    if ser is None:
        return None

    try:
        line = ser.readline().decode().strip()

        if line.startswith("S1:"):
            parts = line.split()
            values = {}

            for p in parts:
                k, v = p.split(":")
                values[k] = float(v)

            return values
    except:
        return None
    

import time

if "auto_refresh" not in st.session_state:
    st.session_state.auto_refresh = True



# ======================================================
# SOC / SOH ESTIMATION
# ======================================================

def calculate_soc(voltage, v_min=3.0, v_max=4.2):
    soc = (voltage - v_min) / (v_max - v_min) * 100
    return float(np.clip(soc, 0, 100))


def calculate_soh(voltage, avg_voltage, temperature):
    delta_t = max(0, temperature - 25)
    delta_v = abs(voltage - avg_voltage)
    soh = 100 - (0.3 * delta_t) - (80 * delta_v)
    return float(np.clip(soh, 80, 100))


def classify_severity(fault):
    if fault["THERMAL"]:
        return "CRITICAL"
    if fault["OV"] or fault["UV"]:
        return "WARNING"
    return "NORMAL"

def auto_refresh(seconds=1):
    if "last_refresh" not in st.session_state:
        st.session_state.last_refresh = time.time()

    if time.time() - st.session_state.last_refresh > seconds:
        st.session_state.last_refresh = time.time()
        st.rerun()


# ======================================================
# MOCK BMS DATA
# ======================================================

def generate_cell_data():
    cells = []
    series_voltages = np.random.normal(3.80, 0.02, N_CELLS_SERIES)
    temperature = np.random.normal(39.5, 1.0)
    avg_voltage = np.mean(series_voltages)

    for i, v in enumerate(series_voltages):
        soc = calculate_soc(v)
        soh = calculate_soh(v, avg_voltage, temperature)

        fault = {
            "OV": v > 4.2,
            "UV": v < 3.0,
            "THERMAL": temperature > 45
        }

        cells.append({
            "series_id": f"S{i+1}",
            "voltage": round(v, 3),
            "temperature": round(temperature, 1),
            "SOC": round(soc, 1),
            "SOH": round(soh, 1),
            "hotspot_prob": round(np.clip((temperature - 30) * 3, 0, 100), 1),
            "fault": fault,
            "severity": classify_severity(fault)
        })

    return cells


def generate_timeseries(days=3):
    time_index = pd.date_range(
        datetime.now() - timedelta(days=days),
        periods=days * 24,
        freq="H"
    )

    voltage = np.linspace(3.95, 3.70, len(time_index)) + np.random.normal(0, 0.02, len(time_index))
    temperature = np.random.normal(39, 1.5, len(time_index))

    soc = [calculate_soc(v) for v in voltage]
    avg_v = np.mean(voltage)
    soh = [calculate_soh(v, avg_v, t) for v, t in zip(voltage, temperature)]

    return pd.DataFrame({
        "time": time_index,
        "voltage": voltage,
        "temperature": temperature,
        "SOC": soc,
        "SOH": soh
    })

# ======================================================
# SENSOR + GRID LOGIC (PHYSICALLY CONSISTENT)
# ======================================================

def mock_sensor_data():
    return {
        "S1": np.random.normal(38, 1),
        "S2": np.random.normal(39, 1),
        "S3": np.random.normal(40, 1),
        "S4": np.random.normal(39.5, 1)
    }


def mock_cell_grid_from_sensors(sensors):
    """
    Generate 3x3 grid from sensor values
    Force ONE cell to be a hotspot (demo purpose)
    """
    base = list(sensors.values())
    grid = []

    for i in range(3):
        row = []
        for j in range(3):
            temp = np.random.normal(base[i % len(base)], 0.6)
            row.append({
                "state": "O",
                "temp": round(temp, 1)
            })
        grid.append(row)

    # 🔴 FORCE ONE CELL TO BE HOT (example: middle cell)
    grid[1][1]["state"] = "X"
    grid[1][1]["temp"] = 45.8   # clearly above threshold

    return grid


def render_cell_grid(grid):
    html = "<table style='border-collapse: collapse; margin:auto;'>"
    for row in grid:
        html += "<tr>"
        for cell in row:
            color = "#2ecc71" if cell["state"] == "O" else "#e74c3c"
            html += f"""
            <td style="
                width:70px;
                height:70px;
                text-align:center;
                border:2px solid black;
                background:{color};
                color:white;
                font-size:16px;
                font-weight:bold;">
                {cell['state']}<br>{cell['temp']}°C
            </td>
            """
        html += "</tr>"
    html += "</table>"
    return html
df = pd.read_csv("bms_raw_dataset_with_anomalies.csv")

df["timestamp"] = pd.to_datetime(df["timestamp"])
df = df.sort_values("timestamp")

# temperature
df["temperature_C"] = df[["S1_temp","S2_temp","S3_temp","S4_temp"]].mean(axis=1)

# voltage features
cell_cols = ["Cell1_V","Cell2_V","Cell3_V","Cell4_V","Cell5_V"]

df["avg_cell_voltage"] = df[cell_cols].mean(axis=1)
df["max_cell_voltage"] = df[cell_cols].max(axis=1)
df["min_cell_voltage"] = df[cell_cols].min(axis=1)

df["voltage_imbalance"] = df["max_cell_voltage"] - df["min_cell_voltage"]

# SOC estimation
df["SOC"] = ((df["avg_cell_voltage"] - 3.3) / (4.2 - 3.3)) * 100
df["SOC"] = df["SOC"].clip(0,100)

# degradation model
deg = 0.02
degradation = []

for i,row in df.iterrows():

    stress = (
        0.4 * (row["temperature_C"]/50) +
        0.4 * (row["voltage_imbalance"]/0.05) +
        0.2 * (abs(row["Pack_Current"])/10)
    )

    deg += stress * 0.0001
    deg = min(deg,1)

    degradation.append(deg)

df["degradation_index"] = degradation
df["SOH"] = 100 * (1 - df["degradation_index"])

df["RUL_cycles"] = 1000 * (1 - df["degradation_index"])
df["Remaining_Days"] = df["RUL_cycles"] / 2

latest = df.iloc[-1]

pack_soc = latest["SOC"]
current_soh = latest["SOH"]
pack_temp = latest["temperature_C"]
pack_voltage = latest["Pack_Voltage"]
pack_current = latest["Pack_Current"]
# ======================================================
# GLOBAL STATE
# ======================================================

# ======================================================
# HEADER
# ======================================================

st.title("🔋 AI-Powered Energy-Efficient Storage System")

# latest values from dataset
latest = df.iloc[-1]

pack_soc = round(latest["SOC"],2)
pack_soh = round(latest["SOH"],2)
pack_temp = round(latest["temperature_C"],2)

# system status logic
if pack_temp > 45:
    system_state = "CRITICAL"
    status_color = "red"
elif pack_temp > 40:
    system_state = "WARNING"
    status_color = "orange"
else:
    system_state = "NORMAL"
    status_color = "green"

# ===============================
# SYSTEM SUMMARY PANEL
# ===============================

col1,col2,col3,col4,col5 = st.columns(5)

col1.markdown(f"""
**Battery ID**

{BATTERY_ID}
""")

col2.metric(
    "State of Charge",
    f"{pack_soc}%"
)

col3.metric(
    "State of Health",
    f"{pack_soh}%"
)

col4.metric(
    "Pack Temperature",
    f"{pack_temp} °C"
)

col5.markdown(f"""
**System Status**

<span style='color:{status_color}; font-size:22px; font-weight:bold'>
{system_state}
</span>
""", unsafe_allow_html=True)

st.divider()


# ==========================================================
# GLOBAL DATA LOADING (MUST BE ABOVE TABS)

# ==========================================================
# GLOBAL METRICS


# ======================================================
# TABS
# ======================================================

tab1, tab2, tab3 = st.tabs([
    
    "📉 AI-Powered Degradation & RUL Intelligence",
    "🧬 Digital Twin",
    "⚡ Grid Dispatch"
])

# ======================================================
# GLOBAL DATA ENGINE
# ======================================================





# ======================================================
# TAB 2 — DEGRADATION
# ======================================================

with tab1:

    import numpy as np
    import plotly.graph_objects as go
    import streamlit as st

    st.subheader("📉 AI-Powered Degradation & RUL Intelligence")

    # =====================================================
    # INITIALIZE PERSISTENT DEGRADATION
    # =====================================================

    if "degradation_index" not in st.session_state:
        st.session_state.degradation_index = 0.02

    # =====================================================
    # USE LATEST DATA FROM GLOBAL DATAFRAME
    # =====================================================

    latest = df.iloc[-1]

    temperature = latest["temperature_C"]
    current = abs(latest["Pack_Current"])
    imbalance = latest["voltage_imbalance"]
    soc = latest["SOC"]

    # =====================================================
    # LIVE DEGRADATION MODEL
    # =====================================================

    temp_factor = temperature / 50
    imbalance_factor = imbalance / 0.05
    current_factor = current / 10

    stress = (
        0.45 * temp_factor +
        0.35 * imbalance_factor +
        0.20 * current_factor
    )

        # slow degradation for demo stability
    deg_growth = stress * 0.00002

    st.session_state.degradation_index += deg_growth
    st.session_state.degradation_index = min(st.session_state.degradation_index, 1)

    degradation_index = st.session_state.degradation_index

    # =====================================================
    # SOH + RUL CALCULATIONS
    # =====================================================

    SOH = 100 * (1 - degradation_index)

    total_cycles_life = 1000

    RUL_cycles = total_cycles_life * (1 - degradation_index)

    Remaining_Days = RUL_cycles / 2

    # =====================================================
    # BATTERY HEALTH OVERVIEW
    # =====================================================

    st.markdown("### Battery Health Overview")

    c1, c2, c3 = st.columns(3)

    # SOC GAUGE
    with c1:

        fig_soc = go.Figure(go.Indicator(
            mode="gauge+number",
            value=soc,
            title={'text': "State of Charge (%)"},
            gauge={
                'axis': {'range': [0, 100]},
                'steps': [
                    {'range': [0, 30], 'color': "#ff4b4b"},
                    {'range': [30, 70], 'color': "#ffa500"},
                    {'range': [70, 100], 'color': "#00ff9c"}
                ]
            }
        ))

        fig_soc.update_layout(height=300)

        st.plotly_chart(fig_soc, use_container_width=True)

    # SOH GAUGE
    with c2:

        fig_soh = go.Figure(go.Indicator(
            mode="gauge+number",
            value=SOH,
            title={'text': "State of Health (%)"},
            gauge={
                'axis': {'range': [70, 100]},
                'steps': [
                    {'range': [70, 80], 'color': "#ff4b4b"},
                    {'range': [80, 90], 'color': "#ffa500"},
                    {'range': [90, 100], 'color': "#00ff9c"}
                ]
            }
        ))

        fig_soh.update_layout(height=300)

        st.plotly_chart(fig_soh, use_container_width=True)

    # DEGRADATION INDICATOR
    with c3:

        fig_deg = go.Figure(go.Indicator(
            mode="number",
            value=degradation_index,
            number={'valueformat': '.4f'},
            title={'text': "Degradation Index"}
        ))

        fig_deg.update_layout(height=300)

        st.plotly_chart(fig_deg, use_container_width=True)

    st.divider()

    # =====================================================
    # LIFETIME FORECAST
    # =====================================================

    st.markdown("### Battery Lifetime Forecast")

    r1, r2 = st.columns(2)

    # Previous values for delta display
    if "prev_rul" not in st.session_state:
        st.session_state.prev_rul = RUL_cycles

    if "prev_days" not in st.session_state:
        st.session_state.prev_days = Remaining_Days

    with r1:
        st.metric(
            "Remaining Useful Life (Cycles)",
            int(RUL_cycles),
            delta=int(RUL_cycles - st.session_state.prev_rul)
        )

    with r2:
        st.metric(
            "Estimated Remaining Days",
            int(Remaining_Days),
            delta=int(Remaining_Days - st.session_state.prev_days)
        )

    # update previous values
    st.session_state.prev_rul = RUL_cycles
    st.session_state.prev_days = Remaining_Days

    # =============================
    # Lifetime Progress Bar
    # =============================

    life_percent = SOH

    st.progress(life_percent / 100)

    st.caption(f"Battery life remaining: {life_percent:.2f}%")

    st.divider()

    # =====================================================
    # OPERATIONAL STRESS ANALYSIS
    # =====================================================

    st.markdown("### Operational Stress Diagnostics")

    thermal_index = max(0, temperature - 30)

    current_index = current

    imbalance_index = imbalance

    fig_radar = go.Figure()

    fig_radar.add_trace(go.Scatterpolar(
        r=[thermal_index, current_index, imbalance_index],
        theta=["Thermal Stress", "Current Stress", "Voltage Imbalance"],
        fill='toself'
    ))

    fig_radar.update_layout(
        height=400,
        polar=dict(radialaxis=dict(visible=True))
    )

    st.plotly_chart(fig_radar, use_container_width=True)

    st.caption(
        "Higher stress indicators accelerate battery ageing and reduce remaining life."
    )
# ======================================================
# TAB 3 — DIGITAL TWIN
# ======================================================

with tab2:

    import plotly.graph_objects as go
    import plotly.express as px
    import numpy as np

    st.subheader("🧬 Battery Digital Twin")

    latest = df.iloc[-1]

    # =====================================================
    # PACK SNAPSHOT
    # =====================================================

    st.markdown("### Pack Operational Snapshot")

    col1, col2, col3, col4 = st.columns(4)

    pack_voltage = latest["Pack_Voltage"]
    pack_current = latest["Pack_Current"]
    pack_temp = latest["temperature_C"]
    power = pack_voltage * pack_current

    col1.metric("Pack Voltage", f"{pack_voltage:.2f} V")
    col2.metric("Pack Current", f"{pack_current:.2f} A")
    col3.metric("Battery Power", f"{power:.2f} W")
    col4.metric("Pack Temperature", f"{pack_temp:.1f} °C")

    st.divider()

    # =====================================================
    # CELL VOLTAGE DISTRIBUTION
    # =====================================================

    st.markdown("### Cell Voltage Distribution")

    cell_cols = ["Cell1_V","Cell2_V","Cell3_V","Cell4_V","Cell5_V"]
    voltages = latest[cell_cols].values

    fig_bar = px.bar(
        x=cell_cols,
        y=voltages,
        labels={"x":"Cell","y":"Voltage (V)"},
        title="Individual Cell Voltages",
        template="plotly_dark"
    )

    st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()

    # =====================================================
    # THERMAL DIGITAL TWIN MAP
    # =====================================================

    st.markdown("### Thermal Digital Twin")

    temps = [
        latest["S1_temp"],
        latest["S2_temp"],
        latest["S3_temp"],
        latest["S4_temp"]
    ]

    thermal_grid = np.array([
        [temps[0], temps[1]],
        [temps[2], temps[3]]
    ])

    fig_heat = px.imshow(
        thermal_grid,
        text_auto=True,
        color_continuous_scale="Turbo",
        aspect="auto",
        title="Battery Pack Thermal Map (°C)"
    )

    fig_heat.update_layout(template="plotly_dark")

    st.plotly_chart(fig_heat, use_container_width=True)

    st.divider()

    # =====================================================
    # OPERATIONAL STRESS RADAR
    # =====================================================

    st.markdown("### Operational Stress Profile")

    thermal_index = max(0, pack_temp - 30)

    current_index = abs(pack_current)

    imbalance = max(voltages) - min(voltages)

    fig_radar = go.Figure()

    fig_radar.add_trace(go.Scatterpolar(
        r=[thermal_index, current_index, imbalance],
        theta=[
            "Thermal Stress",
            "Current Stress",
            "Voltage Imbalance"
        ],
        fill='toself'
    ))

    fig_radar.update_layout(
        template="plotly_dark",
        polar=dict(radialaxis=dict(visible=True))
    )

    st.plotly_chart(fig_radar, use_container_width=True)

    st.divider()

    # =====================================================
    # SYSTEM RISK ENGINE
    # =====================================================

    st.markdown("### System Stability Index")

    risk_score = int(
        0.45 * thermal_index +
        0.35 * imbalance * 100 +
        0.20 * abs(pack_current)
    )

    risk_score = np.clip(risk_score, 0, 100)

    fig_risk = go.Figure(go.Indicator(
        mode="gauge+number",
        value=risk_score,
        title={'text': "Operational Risk"},
        gauge={
            'axis': {'range': [0,100]},
            'steps': [
                {'range':[0,40],'color':"#00FFAA"},
                {'range':[40,70],'color':"#FFA500"},
                {'range':[70,100],'color':"#FF4B4B"}
            ]
        }
    ))

    fig_risk.update_layout(template="plotly_dark", height=350)

    st.plotly_chart(fig_risk, use_container_width=True)

# ======================================================
# TAB 4 — GRID DISPATCH
# ======================================================
with tab3:

    import plotly.express as px
    import plotly.graph_objects as go
    from docx import Document
    from io import BytesIO
    from groq import Groq
    import numpy as np

    st.subheader("⚡ Energy Dispatch Intelligence")

    latest = df.iloc[-1]

    # =====================================================
    # PACK METRICS
    # =====================================================

    pack_voltage = latest["Pack_Voltage"]
    pack_current = latest["Pack_Current"]

    pack_temp = latest["temperature_C"]

    avg_cell_voltage = latest["avg_cell_voltage"]

    pack_soc = latest["SOC"]
    current_soh = latest["SOH"]

    battery_power = pack_voltage * pack_current

    # =====================================================
    # GRID CONTEXT (SIMULATION)
    # =====================================================

    solar_generation = np.random.uniform(0, 4)  # kW
    grid_price = np.random.choice(["Low", "Medium", "High"])

    demand_load = abs(battery_power) / 1000
    demand_load = round(max(demand_load, 1.5), 2)

    st.markdown("### Energy System Context")

    c1, c2, c3, c4 = st.columns(4)

    c1.metric("Load Demand", f"{demand_load:.2f} kW")
    c2.metric("Solar Generation", f"{solar_generation:.2f} kW")
    c3.metric("Battery SOC", f"{pack_soc:.1f}%")
    c4.metric("Grid Price Signal", grid_price)

    st.divider()

    # =====================================================
    # ENERGY DISPATCH DECISION
    # =====================================================

    st.markdown("### Smart Dispatch Engine")

    battery_score = (
        0.45 * (current_soh / 100) +
        0.35 * (pack_soc / 100) +
        0.20 * (1 - max(pack_temp - 30, 0) / 20)
    )

    solar_score = solar_generation / 4

    grid_score = {"Low":0.9,"Medium":0.6,"High":0.3}[grid_price]

    scores = {
        "Solar": solar_score,
        "Battery": battery_score,
        "Grid": grid_score
    }

    total = sum(scores.values()) + 1e-6

    allocation = {k:v/total*100 for k,v in scores.items()}

    col1, col2 = st.columns(2)

    # ENERGY ALLOCATION PIE
    with col1:

        fig = px.pie(
            names=list(allocation.keys()),
            values=list(allocation.values()),
            hole=0.55,
            template="plotly_dark",
            title="Energy Dispatch Allocation"
        )

        st.plotly_chart(fig, use_container_width=True)

    # POWER FLOW BAR
    with col2:

        fig_bar = px.bar(
            x=list(allocation.keys()),
            y=list(allocation.values()),
            labels={"x":"Energy Source","y":"Contribution (%)"},
            template="plotly_dark",
            title="Power Source Contribution"
        )

        st.plotly_chart(fig_bar, use_container_width=True)

    st.divider()

    # =====================================================
    # OPERATING MODE ENGINE
    # =====================================================

    if pack_temp > 42 or current_soh < 85:

        mode = "Battery Protection Mode"
        description = "Battery usage minimized due to thermal or health constraints."

    elif solar_generation > demand_load:

        mode = "Renewable Priority Mode"
        description = "Solar generation sufficient. Grid import minimized."

    elif grid_price == "Low":

        mode = "Grid Charging Mode"
        description = "Cheap grid electricity available. Charging battery."

    else:

        mode = "Battery Support Mode"
        description = "Battery discharging to support demand."

    st.markdown("### Dispatch Operating Mode")

    st.success(f"**{mode}**")

    st.caption(description)

    st.divider()

    # =====================================================
    # SYSTEM ENERGY FLOW
    # =====================================================

    st.markdown("### System Energy Flow")

    flow_data = {
        "Source": ["Solar", "Battery", "Grid"],
        "Energy Share": list(allocation.values())
    }

    fig_flow = px.bar(
        flow_data,
        x="Source",
        y="Energy Share",
        color="Source",
        template="plotly_dark",
        title="Energy Flow Distribution"
    )

    st.plotly_chart(fig_flow, use_container_width=True)

    st.divider()

    # =====================================================
    # AI REPORT GENERATION
    # =====================================================

    st.subheader("📝 AI Operational Report")

    if "ai_report_buffer" not in st.session_state:
        st.session_state.ai_report_buffer = None
        st.session_state.ai_report_text = None

    if st.button("Generate AI Report"):

        st.session_state.auto_refresh = False

        try:

            client = Groq(api_key="")

            prompt = f"""
            Generate a professional energy storage operational report.

            Metrics:
            Pack Voltage: {pack_voltage:.2f} V
            Pack Current: {pack_current:.2f} A
            Pack Temperature: {pack_temp:.2f} °C
            SOC: {pack_soc:.2f} %
            SOH: {current_soh:.2f} %
            Dispatch Mode: {mode}

            Include:
            - Executive Summary
            - Battery Health Analysis
            - Thermal Risk Assessment
            - Dispatch Optimization
            - Technical Recommendations
            """

            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role":"system","content":"You are a senior energy storage engineer."},
                    {"role":"user","content":prompt}
                ],
                temperature=0.3
            )

            ai_text = response.choices[0].message.content

            st.session_state.ai_report_text = ai_text

            doc = Document()
            doc.add_heading("Energy Storage Operational Report", level=1)
            doc.add_paragraph(ai_text)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.session_state.ai_report_buffer = buffer

            st.success("AI Report Generated")

        except Exception as e:

            st.error(f"Groq API Error: {e}")

        finally:

            st.session_state.auto_refresh = True

    if st.session_state.ai_report_text:

        st.markdown("### AI Insights")
        st.write(st.session_state.ai_report_text)

    if st.session_state.ai_report_buffer:

        st.download_button(
            label="⬇ Download Report (.docx)",
            data=st.session_state.ai_report_buffer,
            file_name="Energy_Storage_AI_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# ======================================================
# AUTO REFRESH ENGINE (MUST BE LAST)
# ======================================================

if st.session_state.auto_refresh:
    time.sleep(1)
    st.rerun()